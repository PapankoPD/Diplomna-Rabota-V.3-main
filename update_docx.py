import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

def update_docx(docx_path, text_path):
    print(f"Loading document {docx_path}...")
    doc = Document(docx_path)
    
    print(f"Reading content from {text_path}...")
    with open(text_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    print("Appending content...")
    # Add a blank line before appending if desired
    doc.add_paragraph("")
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run = p.add_run(line)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # If it's a heading (starts with a number like 4. or 4.1.) we can make it bold
        if line.startswith("4. Анализ") or line.startswith("4.1.") or line.startswith("4.2.") or line.startswith("4.3.") or line.startswith("4.4."):
            run.bold = True
            
    doc.save(docx_path)
    print("Successfully updated the document.")

if __name__ == "__main__":
    update_docx("2.docx", "section9.txt")
