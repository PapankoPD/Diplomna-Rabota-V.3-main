import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def add_use_case_table(doc, index, title, actor, precondition, main_scenario, alt_scenario, postcondition, ucid):
    # Insert table
    table = doc.add_table(rows=6, cols=2, style='Table Grid')
    
    # Header format
    headers = ["Идентификатор", "Име", "Участници (Актьори)", "Предусловие", "Основен сценарий", "Постусловие"]
    data = [ucid, title, actor, precondition, main_scenario, postcondition]
    
    for i in range(6):
        # Set bold on first col
        cell_left = table.cell(i, 0)
        cell_left.text = headers[i]
        for p in cell_left.paragraphs:
            for run in p.runs:
                run.font.bold = True
                
        # Set text on right col
        cell_right = table.cell(i, 1)
        cell_right.text = data[i]
        
    # Make left col smaller
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)
        
    return table

def move_element_after(target_element, new_element):
    """Moves new_element to appear after target_element in the document XML."""
    target_element.addnext(new_element)

def main():
    doc_path = r'c:\Users\Nik\Desktop\Diplomna-Rabota-V.3-main\2.docx'
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error opening doc: {e}")
        return

    # Find the "Сценарии на използване (Use Cases)" header
    target_p = None
    for p in doc.paragraphs:
        if "Сценарии на използване (Use Cases)" in p.text:
            target_p = p
            break
            
    if not target_p:
        print("Header not found")
        return

    use_cases = [
        ("UC-01", "Вход в системата (Автентикация)", "Потребител (Ученик, Преподавател, Администратор)", "Потребителят е регистриран и има валиден акаунт в системата.", "1. Потребителят въвежда потребителско име и парола.\n2. Системата валидира данните.\n3. Системата генерира JWT токен.\n4. Потребителят е пренасочен към главното табло.\n(Алтернативно: при грешни данни се показва съобщение)", "Потребителят е успешно автентикиран."),
        ("UC-02", "Качване на учебни материали", "Преподавател, Администратор", "Потребителят е логнат и има права за качване (Teacher/Admin).", "1. Преподавателят избира 'Качване на материал'.\n2. Попълва метаданни и прикачва файлове.\n3. Определя права за достъп (кръг на видимост).\n4. Системата запазва файловете и базата данни.\n5. Системата връща визуално потвърждение.", "Материалът е качен и наличен за предвидената група ученици."),
        ("UC-03", "Търсене и филтриране на ресурси", "Всички потребители", "Потребителят е влязъл в системата.", "1. Потребителят въвежда ключова дума.\n2. Системата изпълнява FTS5 заявка.\n3. Връща списък с релевантни материали, съобразен с правата за достъп.", "Потребителят разглежда резултатите от търсенето."),
        ("UC-04", "Управление на архивирани материали", "Администратор / Автор на файла", "Потребителят е логнат и преглежда архивните си ресурси.", "1. Потребителят избира списък 'Архивирани материали'.\n2. Намира конкретен материал.\n3. Избира опция 'Възстанови'.\n4. Системата променя статуса на файла.", "Файлът е възстановен и достъпен в основната библиотека.")
    ]
    
    # We add tables to the end of the doc, then move them right after the header target
    # We must construct them in reverse order so they array correctly when moved repeatedly after `target_p`
    # Or keep track of the last inserted element.
    
    last_inserted_element = target_p._p
    
    for uc in use_cases:
        # Add spacing
        spacer = doc.add_paragraph()
        spacer_elm = spacer._p
        
        table = add_use_case_table(doc, 0, uc[1], uc[2], uc[3], uc[4], "", uc[5], uc[0])
        table_elm = table._tbl
        
        # move spacer
        last_inserted_element.addnext(spacer_elm)
        # move table
        spacer_elm.addnext(table_elm)
        
        last_inserted_element = table_elm

    try:
        doc.save(doc_path)
        print("Success")
    except Exception as e:
        print(f"Error saving: {e}")

if __name__ == '__main__':
    main()
