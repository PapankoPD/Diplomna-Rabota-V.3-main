"""
insert_toc.py
Поправя заглавията в 2.docx и вмъква автоматично Word TOC поле с номера на страници.
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

DOC_PATH = r"c:\Users\Nik\Desktop\Diplomna-Rabota-V.3-main\2.docx"

# ---------------------------------------------------------------------------
# Карта: текст (startswith) -> (Heading ниво, нов текст ако трябва замяна)
# ---------------------------------------------------------------------------
HEADING_MAP = [
    # Ниво 1 – основни глави
    ("Увод (Въведение)",                             1, "3. Увод (Въведение)"),
    ("3. Увод",                                      1, None),
    ("4. Анализ и спецификация",                     1, None),
    ("5. Архитектура на системата",                  1, None),
    ("6. Реализация на софтуерното решение",         1, None),
    ("7. Заключение",                                1, None),
    ("8. Използвана литература",                     1, None),
    ("9. Приложения",                                1, None),

    # Ниво 2 – подточки
    ("4.1.",  2, None),
    ("4.2.",  2, None),
    ("4.3.",  2, None),
    ("4.4.",  2, None),
    ("5.1.",  2, None),
    ("5.2.",  2, None),
    ("5.3.",  2, None),
    ("5.4.",  2, None),
    ("5.5.",  2, None),
    ("5.6.",  2, None),
    ("6.1.",  2, None),
    ("6.2.",  2, None),
    ("6.3.",  2, None),
    ("6.4.",  2, None),
    ("6.5.",  2, None),
    ("7.1.",  2, None),
    ("7.2.",  2, None),
    ("7.3.",  2, None),

    # Ниво 3 – под-под-точки
    ("6.1.1.", 3, None),
    ("6.1.2.", 3, None),
    ("6.1.3.", 3, None),
    ("6.1.4.", 3, None),
    ("6.1.5.", 3, None),
    ("6.3.1.", 3, None),
    ("6.3.2.", 3, None),
    ("6.3.3.", 3, None),
    ("6.3.4.", 3, None),
]

TOC_HEADING_TEXT = "СЪДЪРЖАНИЕ"


def heading_style_name(level: int) -> str:
    return f"Heading {level}"


def make_toc_field_paragraph(doc: Document):
    """Създава параграф с Word TOC поле (автоматично генерирано от Word)."""
    p = OxmlElement("w:p")

    # Paragraph properties – без специален стил
    pPr = OxmlElement("w:pPr")
    p.append(pPr)

    def run_with_fld(fld_type=None, instr=None):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        noProof = OxmlElement("w:noProof")
        rPr.append(noProof)
        r.append(rPr)
        if fld_type:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_type)
            if fld_type == "begin":
                fc.set(qn("w:dirty"), "true")
            r.append(fc)
        if instr:
            it = OxmlElement("w:instrText")
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = instr
            r.append(it)
        return r

    p.append(run_with_fld(fld_type="begin"))
    p.append(run_with_fld(instr=' TOC \\o "1-3" \\h \\z \\u '))
    p.append(run_with_fld(fld_type="separate"))
    p.append(run_with_fld(fld_type="end"))
    return p


def apply_headings(doc: Document):
    """Прилага Heading стилове към параграфи, където текстът съвпада."""
    changed = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        for prefix, level, replacement in HEADING_MAP:
            if text.startswith(prefix):
                style_name = heading_style_name(level)
                try:
                    para.style = doc.styles[style_name]
                except KeyError:
                    print(f"  [WARN] Стил '{style_name}' не е намерен, пропускам.")
                    continue
                if replacement and text == prefix:
                    # Замени текста само ако е точно съвпадение
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = replacement
                    else:
                        para.add_run(replacement)
                changed += 1
                print(f"  ✅ H{level}: {text[:70]}")
                break
    return changed


def replace_toc_section(doc: Document):
    """
    Намира параграфа 'СЪДЪРЖАНИЕ', изтрива всичко след него
    до следващата голяма секция и вмъква Word TOC поле.
    """
    body = doc.element.body
    paras = list(body.iterchildren(qn("w:p")))

    toc_idx = None
    for i, p in enumerate(paras):
        text = "".join(r.text or "" for r in p.iter(qn("w:t"))).strip()
        if text.upper() == TOC_HEADING_TEXT.upper():
            toc_idx = i
            break

    if toc_idx is None:
        print("  [WARN] Параграф 'СЪДЪРЖАНИЕ' не е намерен!")
        return False

    # Намери края на TOC блока – първия параграф с Heading стил след TOC
    end_idx = None
    for i in range(toc_idx + 1, len(paras)):
        p = paras[i]
        pPr = p.find(qn("w:pPr"))
        style_val = ""
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_val = pStyle.get(qn("w:val"), "")
        text = "".join(r.text or "" for r in p.iter(qn("w:t"))).strip()

        # Спираме при намиране на Heading 1 стил или Увод текст
        if style_val.startswith("Heading") or text.startswith("3. Увод") or text.startswith("Увод"):
            end_idx = i
            break

    if end_idx is None:
        end_idx = toc_idx + 20  # fallback

    # Изтрий параграфите между СЪДЪРЖАНИЕ и края
    to_remove = paras[toc_idx + 1 : end_idx]
    print(f"  🗑  Изтривам {len(to_remove)} ръчни TOC реда...")
    for p in to_remove:
        body.remove(p)

    # Вмъкни TOC полето веднага след СЪДЪРЖАНИЕ параграфа
    toc_para = paras[toc_idx]
    toc_field = make_toc_field_paragraph(doc)
    toc_para.addnext(toc_field)
    print("  ✅ Word TOC поле вмъкнато успешно.")
    return True


def main():
    print(f"📄 Отварям: {DOC_PATH}")
    doc = Document(DOC_PATH)

    print("\n🔖 Прилагам Heading стилове...")
    n = apply_headings(doc)
    print(f"   Общо: {n} заглавия маркирани.")

    print("\n📑 Заменям ръчния TOC с Word TOC поле...")
    replace_toc_section(doc)

    doc.save(DOC_PATH)
    print(f"\n✅ Документът е записан: {DOC_PATH}")
    print("\n⚠️  ВАЖНО: Отвори 2.docx в Microsoft Word и натисни:")
    print("   Ctrl+A → F9  (за да се обнови TOC с реалните номера на страници)")


if __name__ == "__main__":
    main()
